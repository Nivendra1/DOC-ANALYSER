import fitz  # PyMuPDF
import docx
import pytesseract
from PIL import Image
from langdetect import detect, LangDetectException


def extract_text(file_path: str, content_type: str) -> dict:
    if content_type == "application/pdf":
        return extract_from_pdf(file_path)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_from_docx(file_path)
    else:
        return extract_from_image(file_path)


def extract_from_pdf(file_path: str) -> dict:
    doc = fitz.open(file_path)
    full_text = []

    for page in doc:
        text = page.get_text("text")
        full_text.append(text)

    combined = "\n".join(full_text).strip()

    # If text layer is empty or very short, use OCR
    if len(combined) < 50:
        combined = ocr_pdf(file_path)
        method = "ocr"
    else:
        method = "native"

    doc.close()
    return build_result(combined, method)


def ocr_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    texts = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        texts.append(text)
    doc.close()
    return "\n".join(texts).strip()


def extract_from_docx(file_path: str) -> dict:
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also extract tables
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    combined = "\n".join(paragraphs).strip()
    return build_result(combined, "native")


def extract_from_image(file_path: str) -> dict:
    img = Image.open(file_path)
    # Improve OCR accuracy
    img = img.convert("L")  # grayscale
    text = pytesseract.image_to_string(img, config="--oem 3 --psm 6")
    return build_result(text.strip(), "ocr")


def detect_language(text: str) -> str:
    try:
        return detect(text[:500])
    except LangDetectException:
        return "en"


def build_result(text: str, method: str) -> dict:
    words = text.split()
    word_count = len(words)
    reading_time = round(word_count / 200, 1)  # avg reading speed

    return {
        "text": text,
        "method": method,
        "word_count": word_count,
        "reading_time_minutes": reading_time,
        "language": detect_language(text)
    }
